import os
import re
import logging
import asyncio
from typing import Tuple, Optional, Dict
from dotenv import load_dotenv
from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes, ConversationHandler
from telegram.error import TimedOut, NetworkError, RetryAfter
import docx
import fitz  # PyMuPDF

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Constants
SUPPORTED_EXTENSIONS = {'.pdf', '.docx'}
MAX_RETRIES = 3
RETRY_DELAY = 1  # seconds

# Conversation states
WAITING_FOR_NAME, WAITING_FOR_ROLLNO = range(2)

# Store user data temporarily
user_data: Dict[int, Dict] = {}

async def send_message_with_retry(update: Update, text: str, max_retries: int = MAX_RETRIES) -> bool:
    """Send a message with retry logic."""
    for attempt in range(max_retries):
        try:
            await update.message.reply_text(text)
            return True
        except (TimedOut, NetworkError) as e:
            if attempt < max_retries - 1:
                logger.warning(f"Attempt {attempt + 1} failed: {str(e)}. Retrying...")
                await asyncio.sleep(RETRY_DELAY)
            else:
                logger.error(f"Failed to send message after {max_retries} attempts: {str(e)}")
                return False
        except RetryAfter as e:
            logger.warning(f"Rate limited. Waiting {e.retry_after} seconds...")
            await asyncio.sleep(e.retry_after)
        except Exception as e:
            logger.error(f"Unexpected error while sending message: {str(e)}")
            return False
    return False

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Send a message when the command /start is issued."""
    welcome_message = (
        "👋 Welcome to the Document Footer Bot!\n\n"
        "To use this bot:\n"
        "1. Send a PDF or DOCX file\n"
        "2. When prompted, send your name\n"
        "3. When prompted, send your roll number\n\n"
        "The bot will add this information as a footer to your document and send it back."
    )
    await send_message_with_retry(update, welcome_message)

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Handle incoming documents and start the conversation."""
    try:
        # Get the document
        document = update.message.document
        file_name = document.file_name
        file_ext = os.path.splitext(file_name)[1].lower()
        
        # Get username or user ID as fallback
        user = update.effective_user
        user_identifier = user.username if user.username else str(user.id)
        
        # Check file type
        if file_ext not in SUPPORTED_EXTENSIONS:
            await send_message_with_retry(update, f"❌ Unsupported file type. Please send a PDF or DOCX file.")
            return ConversationHandler.END
        
        # Download the file
        file = await context.bot.get_file(document.file_id)
        file_path = f"{user_identifier}_{file_name}"
        await file.download_to_drive(file_path)
        
        # Store file information in user_data
        user_id = update.effective_user.id
        user_data[user_id] = {
            'file_path': file_path,
            'file_ext': file_ext,
            'file_name': file_name,
            'user_identifier': user_identifier
        }
        
        # Ask for name
        await send_message_with_retry(update, "Please send your name:")
        return WAITING_FOR_NAME
        
    except Exception as e:
        logger.error(f"Error handling document: {str(e)}")
        await send_message_with_retry(update, "❌ An error occurred while processing your document. Please try again.")
        return ConversationHandler.END

async def handle_name(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Handle the user's name and ask for roll number."""
    user_id = update.effective_user.id
    name = update.message.text.strip()
    
    if not name:
        await send_message_with_retry(update, "❌ Please provide a valid name:")
        return WAITING_FOR_NAME
    
    # Store name in user_data
    user_data[user_id]['name'] = name
    
    # Ask for roll number
    await send_message_with_retry(update, "Please send your roll number:")
    return WAITING_FOR_ROLLNO

async def handle_rollno(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Handle the user's roll number and process the document."""
    user_id = update.effective_user.id
    roll_no = update.message.text.strip()
    
    if not roll_no:
        await send_message_with_retry(update, "❌ Please provide a valid roll number:")
        return WAITING_FOR_ROLLNO
    
    try:
        # Get stored data
        data = user_data[user_id]
        file_path = data['file_path']
        file_ext = data['file_ext']
        name = data['name']
        user_identifier = data['user_identifier']
        
        # Process the document
        if file_ext == '.docx':
            output_path = await process_docx(file_path, name, roll_no, user_identifier)
        else:  # PDF
            output_path = await process_pdf(file_path, name, roll_no, user_identifier)
        
        # Send the processed file with retry logic
        for attempt in range(MAX_RETRIES):
            try:
                await update.message.reply_document(
                    document=open(output_path, 'rb'),
                    filename=os.path.basename(output_path)
                )
                break
            except (TimedOut, NetworkError) as e:
                if attempt < MAX_RETRIES - 1:
                    logger.warning(f"Attempt {attempt + 1} failed: {str(e)}. Retrying...")
                    await asyncio.sleep(RETRY_DELAY)
                else:
                    raise
        
        # Clean up
        if os.path.exists(file_path):
            os.remove(file_path)
        if os.path.exists(output_path):
            os.remove(output_path)
        del user_data[user_id]
        
        await send_message_with_retry(update, "✅ Document processed successfully!")
        return ConversationHandler.END
        
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        await send_message_with_retry(update, "❌ An error occurred while processing your document. Please try again.")
        # Clean up
        if user_id in user_data:
            if os.path.exists(user_data[user_id]['file_path']):
                os.remove(user_data[user_id]['file_path'])
            del user_data[user_id]
        return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Cancel the conversation and clean up."""
    user_id = update.effective_user.id
    if user_id in user_data:
        if os.path.exists(user_data[user_id]['file_path']):
            os.remove(user_data[user_id]['file_path'])
        del user_data[user_id]
    
    await send_message_with_retry(update, "Operation cancelled. Send /start to begin again.")
    return ConversationHandler.END

async def process_docx(file_path: str, name: str, roll_no: str, user_identifier: str) -> str:
    """Process DOCX file and add footer."""
    doc = docx.Document(file_path)
    
    # Add footer to each section
    for section in doc.sections:
        footer = section.footer
        
        # Clear any existing paragraphs
        for paragraph in footer.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
        
        # Create a new paragraph for the footer
        footer_para = footer.add_paragraph()
        
        # Calculate spacing for alignment
        name_text = f"Name: {name}"
        roll_text = f"Roll No: {roll_no}"
        page_text = "Page No:"
        
        # Add name (left)
        name_run = footer_para.add_run(name_text)
        
        # Add spaces for middle alignment (increased spacing)
        footer_para.add_run(" " * 40)
        
        # Add roll number (middle)
        roll_run = footer_para.add_run(roll_text)
        
        # Add spaces for right alignment (increased spacing)
        footer_para.add_run(" " * 40)
        
        # Add page number (right)
        page_run = footer_para.add_run(page_text)
        
        # Set increased font size (12pt)
        for run in footer_para.runs:
            run.font.size = docx.shared.Pt(12)
        
        # Set paragraph alignment to justify
        footer_para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Set paragraph format for better spacing
        paragraph_format = footer_para.paragraph_format
        paragraph_format.space_before = docx.shared.Pt(0)
        paragraph_format.space_after = docx.shared.Pt(0)
        paragraph_format.line_spacing = 1.0
    
    # Save modified document with username_rollno format
    output_path = f"{name}_{roll_no}.docx"
    doc.save(output_path)
    return output_path

async def process_pdf(file_path: str, name: str, roll_no: str, user_identifier: str) -> str:
    """Process PDF file and add footer."""
    doc = fitz.open(file_path)
    
    # Add footer to each page
    for page in doc:
        # Get page dimensions
        rect = page.rect
        
        # Define margins and sizes
        margin = 40  # Margin from edges
        footer_height = 50  # Increased height of footer area
        font_size = 12  # Font size
        
        # Calculate positions for the three parts
        y_pos = rect.height - margin  # Position from bottom
        
        # Calculate text widths for better positioning
        name_width = len(f"Name: {name}") * font_size * 0.5  # Approximate width
        roll_width = len(f"Roll No: {roll_no}") * font_size * 0.5
        
        # Position calculations
        left_pos = margin  # Left align
        center_pos = (rect.width - roll_width) / 2  # Center align
        right_pos = rect.width - margin - 60  # Right align
        
        # Add the three parts of the footer with bold font
        page.insert_text(
            point=(left_pos, y_pos),
            text=f"Name: {name}",
            fontsize=font_size,
            color=(0, 0, 0)
        )
        
        page.insert_text(
            point=(center_pos, y_pos),
            text=f"Roll No: {roll_no}",
            fontsize=font_size,
            color=(0, 0, 0)
        )
        
        page.insert_text(
            point=(right_pos, y_pos),
            text="Page No:",
            fontsize=font_size,
            color=(0, 0, 0)
        )
    
    # Save modified document with username_rollno format
    output_path = f"{name}_{roll_no}.pdf"
    doc.save(output_path)
    doc.close()
    return output_path

def main() -> None:
    """Start the bot."""
    # Get bot token from environment variable
    token = os.getenv('TELEGRAM_BOT_TOKEN')
    if not token:
        logger.error("No token found. Please set TELEGRAM_BOT_TOKEN in .env file")
        return

    # Create the Application
    application = Application.builder().token(token).build()

    # Create conversation handler
    conv_handler = ConversationHandler(
        entry_points=[MessageHandler(filters.Document.ALL, handle_document)],
        states={
            WAITING_FOR_NAME: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_name),
                CommandHandler('cancel', cancel)
            ],
            WAITING_FOR_ROLLNO: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_rollno),
                CommandHandler('cancel', cancel)
            ],
        },
        fallbacks=[CommandHandler('cancel', cancel)]
    )

    # Add handlers
    application.add_handler(CommandHandler("start", start))
    application.add_handler(conv_handler)

    # Start the Bot
    application.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == '__main__':
    main() 